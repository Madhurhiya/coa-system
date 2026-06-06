import os, json
from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse, JsonResponse
from django.db.models import Q
from django.contrib.auth.decorators import login_required
from django.template.loader import get_template
from django.conf import settings
from xhtml2pdf import pisa
from itertools import groupby
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from io import BytesIO

from django.contrib.auth.models import User
from django.contrib import messages
from .models import (Category, TestParameter, COA, COAResult,
                     COACustomField, COALabel, ItemMaster, Customer, UserProfile)
from .forms import COAForm


def link_callback(uri, rel):
    if uri.startswith(settings.STATIC_URL):
        relative = uri[len(settings.STATIC_URL):]
        for static_dir in getattr(settings, 'STATICFILES_DIRS', []):
            path = os.path.join(static_dir, relative)
            if os.path.exists(path):
                return path
        if getattr(settings, 'STATIC_ROOT', None):
            path = os.path.join(settings.STATIC_ROOT, relative)
            if os.path.exists(path):
                return path
    if hasattr(settings, 'MEDIA_URL') and uri.startswith(settings.MEDIA_URL):
        path = os.path.join(settings.MEDIA_ROOT, uri[len(settings.MEDIA_URL):])
        if os.path.exists(path):
            return path
    return uri


@login_required
def item_lookup(request):
    name = request.GET.get('name', '').strip()
    if not name:
        return JsonResponse({})
    try:
        item = ItemMaster.objects.get(item_name__iexact=name)
        return JsonResponse({
            'botanical_name': item.botanical_name,
            'plant_part':     item.plant_part,
            'item_category':  item.item_category,
        })
    except ItemMaster.DoesNotExist:
        return JsonResponse({})


@login_required
def item_search(request):
    q     = request.GET.get('q', '').strip()
    limit = int(request.GET.get('limit', 20))
    if not q:
        return JsonResponse({'results': []})
    items = ItemMaster.objects.filter(
        item_name__icontains=q
    ).values('item_name', 'item_category', 'botanical_name', 'plant_part')[:limit]
    return JsonResponse({'results': list(items)})


@login_required
def customer_search(request):
    q     = request.GET.get('q', '').strip()
    limit = int(request.GET.get('limit', 20))
    if not q:
        return JsonResponse({'results': []})
    customers = Customer.objects.filter(
        name__icontains=q
    ).values_list('name', flat=True)[:limit]
    return JsonResponse({'results': list(customers)})


@login_required
def coa_list(request):
    query    = request.GET.get('q', '').strip()
    customer = request.GET.get('customer', '').strip()
    item     = request.GET.get('item', '').strip()
    coas     = COA.objects.all().order_by('-created_at')

    if query:
        coas = coas.filter(
            Q(product_name__icontains=query)  |
            Q(batch_no__icontains=query)       |
            Q(customer_name__icontains=query)  |
            Q(botanical_name__icontains=query)
        )
    if customer:
        coas = coas.filter(customer_name__icontains=customer)
    if item:
        coas = coas.filter(product_name__icontains=item)

    return render(request, 'coa/coa_list.html', {
        'coas':     coas,
        'query':    query,
        'customer': customer,
        'item':     item,
    })


@login_required
def delete_coa(request, coa_id):
    coa = get_object_or_404(COA, id=coa_id)
    if request.method == 'POST':
        coa.delete()
        return redirect('coa_list')
    return render(request, 'coa/confirm_delete.html', {'coa': coa})


@login_required
def create_coa(request):
    form = COAForm(request.POST or None)
    grouped_parameters = []
    is_dry_extract = False

    if request.method == 'POST':

        if 'load_fields' in request.POST:
            category_id = request.POST.get('category')
            if category_id:
                parameters = TestParameter.objects.filter(
                    category_id=category_id
                ).select_related('group').order_by('group__order', 'order')
                for group_key, group_items in groupby(parameters, key=lambda p: p.group):
                    grouped_parameters.append({'group': group_key, 'params': list(group_items)})
                try:
                    is_dry_extract = Category.objects.get(id=category_id).is_dry_extract()
                except:
                    pass

        elif 'save_coa' in request.POST:
            if form.is_valid():
                coa = form.save(commit=False)
                coa.created_by = request.user
                coa.save()

                category_id = request.POST.get('category')
                parameters  = TestParameter.objects.filter(category_id=category_id)

                for param in parameters:
                    value             = request.POST.get(f'param_{param.id}', '').strip()
                    standard_override = request.POST.get(f'standard_{param.id}', '').strip()
                    reference         = request.POST.get(f'reference_{param.id}', '').strip()
                    if value:
                        COAResult.objects.create(
                            coa=coa, parameter=param, result=value,
                            reference=reference,
                            standard_override=standard_override if standard_override != param.specification else ''
                        )

                names      = request.POST.getlist('custom_field_name')
                specs      = request.POST.getlist('custom_field_spec')
                results    = request.POST.getlist('custom_field_result')
                references = request.POST.getlist('custom_field_reference')
                headings   = request.POST.getlist('custom_is_heading')
                for i, name in enumerate(names):
                    name = name.strip()
                    if name:
                        is_heading = (headings[i] == '1') if i < len(headings) else False
                        COACustomField.objects.create(
                            coa=coa, field_name=name,
                            specification=specs[i]       if i < len(specs)       else '',
                            result=results[i]             if i < len(results)     else '',
                            reference=references[i]       if i < len(references)  else '',
                            order=i,
                            is_heading=is_heading,
                        )
                return redirect('coa_detail', coa_id=coa.id)

            else:
                category_id = request.POST.get('category')
                if category_id:
                    parameters = TestParameter.objects.filter(
                        category_id=category_id
                    ).select_related('group').order_by('group__order', 'order')
                    for group_key, group_items in groupby(parameters, key=lambda p: p.group):
                        grouped_parameters.append({'group': group_key, 'params': list(group_items)})

    return render(request, 'coa/create_coa.html', {
        'form':               form,
        'grouped_parameters': grouped_parameters,
        'posted_data':        request.POST if request.method == 'POST' else {},
        'prev_results':       {},
        'prev_custom':        [],
        'is_clone':           False,
        'is_dry_extract':     is_dry_extract,
    })


@login_required
def clone_coa(request, coa_id):
    original = get_object_or_404(COA, id=coa_id)

    parameters = TestParameter.objects.filter(
        category=original.category
    ).select_related('group').order_by('group__order', 'order')

    grouped_parameters = []
    for group_key, group_items in groupby(parameters, key=lambda p: p.group):
        grouped_parameters.append({'group': group_key, 'params': list(group_items)})

    prev_results = {
        r.parameter_id: {
            'result':    r.result,
            'reference': r.reference,
            'standard':  r.standard_override or r.parameter.specification,
        }
        for r in original.results.all().select_related('parameter')
    }

    prev_custom = list(original.custom_fields.values(
        'field_name', 'specification', 'result', 'reference', 'is_heading'
    ))

    if request.method == 'POST':
        form = COAForm(request.POST)
        if 'save_coa' in request.POST and form.is_valid():
            coa = form.save(commit=False)
            coa.created_by = request.user
            coa.save()

            category_id = request.POST.get('category')
            parameters  = TestParameter.objects.filter(category_id=category_id)

            for param in parameters:
                value             = request.POST.get(f'param_{param.id}', '').strip()
                standard_override = request.POST.get(f'standard_{param.id}', '').strip()
                reference         = request.POST.get(f'reference_{param.id}', '').strip()
                if value:
                    COAResult.objects.create(
                        coa=coa, parameter=param, result=value,
                        reference=reference,
                        standard_override=standard_override if standard_override != param.specification else ''
                    )

            names      = request.POST.getlist('custom_field_name')
            specs      = request.POST.getlist('custom_field_spec')
            results    = request.POST.getlist('custom_field_result')
            references = request.POST.getlist('custom_field_reference')
            headings   = request.POST.getlist('custom_is_heading')
            for i, name in enumerate(names):
                name = name.strip()
                if name:
                    is_heading = (headings[i] == '1') if i < len(headings) else False
                    COACustomField.objects.create(
                        coa=coa, field_name=name,
                        specification=specs[i]       if i < len(specs)       else '',
                        result=results[i]             if i < len(results)     else '',
                        reference=references[i]       if i < len(references)  else '',
                        order=i,
                        is_heading=is_heading,
                    )
            return redirect('coa_detail', coa_id=coa.id)
    else:
        form = COAForm(initial={
            'product_name':   original.product_name,
            'category':       original.category,
            'botanical_name': original.botanical_name,
            'plant_part':     original.plant_part,
            'customer_name':  original.customer_name,
        })

    return render(request, 'coa/create_coa.html', {
        'form':               form,
        'grouped_parameters': grouped_parameters,
        'posted_data':        {},
        'prev_results':       prev_results,
        'prev_custom':        prev_custom,
        'is_clone':           True,
        'original':           original,
    })


@login_required
def coa_detail(request, coa_id):
    coa           = get_object_or_404(COA, id=coa_id)
    results       = coa.results.all().select_related(
        'parameter', 'parameter__group'
    ).order_by('parameter__group__order', 'parameter__order')
    custom_fields = coa.custom_fields.all()
    return render(request, 'coa/coa_detail.html', {
        'coa':           coa,
        'results':       results,
        'custom_fields': custom_fields,
    })


@login_required
def download_coa_pdf(request, coa_id):
    coa     = get_object_or_404(COA, id=coa_id)
    results = coa.results.all().select_related(
        'parameter', 'parameter__group'
    ).order_by('parameter__group__order', 'parameter__order')

    grouped_results = []
    for group_key, group_items in groupby(results, key=lambda r: r.parameter.group):
        grouped_results.append({'group': group_key, 'items': list(group_items)})

    custom_fields   = list(coa.custom_fields.all())
    is_dry_extract  = coa.category.is_dry_extract()

    STATIC = settings.STATIC_URL
    def s(f): return f"{STATIC}images/{f}"

    template = get_template('coa/coa_pdf.html')
    html = template.render({
        'coa':             coa,
        'grouped_results': grouped_results,
        'custom_fields':   custom_fields,
        'is_dry_extract':  is_dry_extract,
        'logo_url':        s('logo.png'),
        'halal_badge_url': s('halal_badge.png'),
        'iso_badge_url':   s('iso_badge.png'),
        'gmp_badge_url':   s('gmp_badge.png'),
        'stamp_url':       s('stamp.png'),
    })

    response = HttpResponse(content_type='application/pdf')
    safe_batch    = coa.batch_no.replace('/', '-')
    safe_product  = coa.product_name.replace(' ', '_').replace('/', '-')
    safe_customer = (coa.customer_name or 'General').replace(' ', '_').replace('/', '-')
    response['Content-Disposition'] = f'attachment; filename="COA_{safe_product}_{safe_customer}_{safe_batch}.pdf"'

    pisa_status = pisa.CreatePDF(html, dest=response, link_callback=link_callback)
    if pisa_status.err:
        return HttpResponse(f'<h2>PDF Error: {pisa_status.err}</h2>', status=500)
    return response


@login_required
def generate_label(request, coa_id):
    coa = get_object_or_404(COA, id=coa_id)
    try:
        label = coa.label
    except Exception:
        label = None

    if request.method == 'POST':
        data = {
            'invoice_no':   request.POST.get('invoice_no', '').strip(),
            'gross_weight': request.POST.get('gross_weight', '').strip(),
            'tare_weight':  request.POST.get('tare_weight', '').strip(),
            'net_weight':   request.POST.get('net_weight', '').strip(),
        }
        if label:
            for k, v in data.items():
                setattr(label, k, v)
            label.save()
        else:
            label = COALabel.objects.create(coa=coa, **data)
        return redirect('download_label_pdf', coa_id=coa.id)

    return render(request, 'coa/generate_label.html', {'coa': coa, 'label': label})


@login_required
def download_label_pdf(request, coa_id):
    coa = get_object_or_404(COA, id=coa_id)
    try:
        label = coa.label
    except Exception:
        return redirect('generate_label', coa_id=coa_id)

    STATIC = settings.STATIC_URL
    def s(f): return f"{STATIC}images/{f}"

    template = get_template('coa/label_pdf.html')
    html = template.render({
        'coa':             coa,
        'label':           label,
        'logo_url':        s('logo.png'),
        'halal_badge_url': s('halal_badge.png'),
        'iso_badge_url':   s('iso_badge.png'),
        'gmp_badge_url':   s('gmp_badge.png'),
    })

    response = HttpResponse(content_type='application/pdf')
    safe_batch   = coa.batch_no.replace('/', '-')
    safe_product = coa.product_name.replace(' ', '_').replace('/', '-')
    response['Content-Disposition'] = f'attachment; filename="LABEL_{safe_batch}_{safe_product}.pdf"'

    pisa_status = pisa.CreatePDF(html, dest=response, link_callback=link_callback)
    if pisa_status.err:
        return HttpResponse(f'<h2>Label PDF Error: {pisa_status.err}</h2>', status=500)
    return response


@login_required
def edit_coa(request, coa_id):
    coa = get_object_or_404(COA, id=coa_id)

    existing_results = {
        r.parameter_id: {
            'result':    r.result or '',
            'reference': r.reference or '',
            'standard':  r.standard_override or r.parameter.specification or '',
            'obj':       r,
        }
        for r in coa.results.all().select_related('parameter')
    }

    parameters = TestParameter.objects.filter(
        category=coa.category
    ).select_related('group').order_by('group__order', 'order')

    grouped_parameters = []
    for group_key, group_items in groupby(parameters, key=lambda p: p.group):
        grouped_parameters.append({'group': group_key, 'params': list(group_items)})

    custom_fields  = list(coa.custom_fields.all())
    is_dry_extract = coa.category.is_dry_extract()

    if request.method == 'POST':
        coa.product_name   = request.POST.get('product_name', coa.product_name).strip()
        coa.botanical_name = request.POST.get('botanical_name', '').strip()
        coa.plant_part     = request.POST.get('plant_part', '').strip()
        coa.customer_name  = request.POST.get('customer_name', '').strip()

        mfg_date_str = request.POST.get('manufacturing_date', '')
        if mfg_date_str:
            from datetime import datetime
            try:
                coa.manufacturing_date = datetime.strptime(mfg_date_str, '%Y-%m-%d').date()
                coa.expiry_date = None
            except ValueError:
                pass
        coa.save()

        for param in TestParameter.objects.filter(category=coa.category):
            value     = request.POST.get(f'param_{param.id}', '').strip()
            standard  = request.POST.get(f'standard_{param.id}', '').strip()
            reference = request.POST.get(f'reference_{param.id}', '').strip()

            if param.id in existing_results:
                r = existing_results[param.id]['obj']
                r.result            = value
                r.reference         = reference
                r.standard_override = standard if standard != param.specification else ''
                r.save()
            elif value:
                COAResult.objects.create(
                    coa=coa, parameter=param, result=value,
                    reference=reference,
                    standard_override=standard if standard != param.specification else ''
                )

        coa.custom_fields.all().delete()
        names      = request.POST.getlist('custom_field_name')
        specs      = request.POST.getlist('custom_field_spec')
        results    = request.POST.getlist('custom_field_result')
        references = request.POST.getlist('custom_field_reference')
        headings   = request.POST.getlist('custom_is_heading')
        for i, name in enumerate(names):
            name = name.strip()
            if name:
                is_heading = (headings[i] == '1') if i < len(headings) else False
                COACustomField.objects.create(
                    coa=coa, field_name=name,
                    specification=specs[i].strip()       if i < len(specs)       else '',
                    result=results[i].strip()             if i < len(results)     else '',
                    reference=references[i].strip()       if i < len(references)  else '',
                    order=i,
                    is_heading=is_heading,
                )

        return redirect('coa_detail', coa_id=coa.id)

    return render(request, 'coa/edit_coa.html', {
        'coa':                coa,
        'grouped_parameters': grouped_parameters,
        'existing_results':   existing_results,
        'custom_fields':      custom_fields,
        'is_dry_extract':     is_dry_extract,
    })


@login_required
def check_result(request):
    import re
    standard = request.POST.get('standard', '').strip()
    result   = request.POST.get('result', '').strip()

    if not standard or not result:
        return JsonResponse({'status': 'unknown'})

    result_nums = re.findall(r'[-+]?\d+\.?\d*', result)
    if not result_nums:
        return JsonResponse({'status': 'unknown', 'message': 'Non-numeric result'})

    result_val  = float(result_nums[0])
    range_match = re.findall(r'[-+]?\d+\.?\d*', standard)

    if len(range_match) >= 2:
        lo, hi = float(range_match[0]), float(range_match[-1])
        if lo > hi:
            lo, hi = hi, lo
        if lo <= result_val <= hi:
            return JsonResponse({'status': 'pass', 'message': f'Within range {lo}-{hi}'})
        else:
            return JsonResponse({'status': 'fail',
                'message': f'Out of range! Expected {lo}-{hi}, got {result_val}'})

    max_match = re.search(r'(?:max|nmt|not more than|<=|<)\s*(\d+\.?\d*)', standard, re.I)
    if max_match:
        limit = float(max_match.group(1))
        if result_val <= limit:
            return JsonResponse({'status': 'pass'})
        return JsonResponse({'status': 'fail', 'message': f'Exceeds max {limit}!'})

    min_match = re.search(r'(?:min|nlt|not less than|>=|>)\s*(\d+\.?\d*)', standard, re.I)
    if min_match:
        limit = float(min_match.group(1))
        if result_val >= limit:
            return JsonResponse({'status': 'pass'})
        return JsonResponse({'status': 'fail', 'message': f'Below minimum {limit}!'})

    return JsonResponse({'status': 'unknown', 'message': 'Could not parse standard'})


@login_required
def product_standards(request):
    name = request.GET.get('name', '').strip()
    if not name:
        return JsonResponse({'standards': {}})
    from .models import ProductStandard
    try:
        ps = ProductStandard.objects.get(product_name__iexact=name)
        return JsonResponse({'standards': ps.get_standards(), 'found': True})
    except ProductStandard.DoesNotExist:
        qs = ProductStandard.objects.filter(product_name__icontains=name).first()
        if qs:
            return JsonResponse({'standards': qs.get_standards(), 'found': True, 'matched': qs.product_name})
        return JsonResponse({'standards': {}, 'found': False})


@login_required
def standards_search(request):
    from .models import ProductStandard
    q     = request.GET.get('q', '').strip()
    limit = int(request.GET.get('limit', 20))
    if not q:
        return JsonResponse({'results': []})
    results = ProductStandard.objects.filter(
        product_name__icontains=q
    ).values_list('product_name', flat=True)[:limit]
    return JsonResponse({'results': list(results)})


@login_required
def old_coa_search(request):
    from .models import OldCOA
    query    = request.GET.get('q', '').strip()
    customer = request.GET.get('customer', '').strip()
    item     = request.GET.get('item', '').strip()
    results  = []

    qs = OldCOA.objects.all()
    if query:
        qs = qs.filter(
            Q(file_name__icontains=query) |
            Q(product__icontains=query)   |
            Q(customer__icontains=query)  |
            Q(batch__icontains=query)
        )
    if customer:
        qs = qs.filter(customer__icontains=customer)
    if item:
        qs = qs.filter(product__icontains=item)

    if query or customer or item:
        results = qs.order_by('file_name')[:100]

    return render(request, 'coa/old_coa_search.html', {
        'query':    query,
        'customer': customer,
        'item':     item,
        'results':  results,
    })


@login_required
def old_coa_detail(request, old_id):
    from .models import OldCOA
    old_coa = get_object_or_404(OldCOA, id=old_id)
    return render(request, 'coa/old_coa_detail.html', {
        'old_coa': old_coa,
        'fields':  old_coa.get_fields(),
    })


@login_required
def clone_from_old(request, old_id):
    from .models import OldCOA
    old_coa    = get_object_or_404(OldCOA, id=old_id)
    old_fields = old_coa.get_fields()
    categories = Category.objects.all().order_by('name')

    if request.method == 'POST' and 'save_coa' in request.POST:
        form = COAForm(request.POST)
        if form.is_valid():
            coa = form.save(commit=False)
            coa.created_by = request.user
            coa.save()

            category_id = request.POST.get('category')
            if category_id:
                for param in TestParameter.objects.filter(category_id=category_id):
                    value     = request.POST.get(f'param_{param.id}', '').strip()
                    standard  = request.POST.get(f'standard_{param.id}', '').strip()
                    reference = request.POST.get(f'reference_{param.id}', '').strip()
                    if value:
                        COAResult.objects.create(
                            coa=coa, parameter=param, result=value,
                            reference=reference,
                            standard_override=standard if standard != param.specification else ''
                        )

            names      = request.POST.getlist('custom_field_name')
            specs      = request.POST.getlist('custom_field_spec')
            results    = request.POST.getlist('custom_field_result')
            references = request.POST.getlist('custom_field_reference')
            headings   = request.POST.getlist('custom_is_heading')
            for i, name in enumerate(names):
                name = name.strip()
                if name:
                    is_heading = (headings[i] == '1') if i < len(headings) else False
                    COACustomField.objects.create(
                        coa=coa, field_name=name,
                        specification=specs[i].strip()       if i < len(specs)       else '',
                        result=results[i].strip()             if i < len(results)     else '',
                        reference=references[i].strip()       if i < len(references)  else '',
                        order=i,
                        is_heading=is_heading,
                    )
            return redirect('coa_detail', coa_id=coa.id)

    else:
        form = COAForm(initial={
            'product_name':   old_coa.product,
            'botanical_name': old_coa.botanical,
            'plant_part':     old_coa.part_used,
            'customer_name':  old_coa.customer,
        })

    grouped_parameters = []
    prev_results       = {}
    selected_cat_id    = request.POST.get('category') if request.method == 'POST' else ''

    if 'load_fields' in request.POST:
        selected_cat_id = request.POST.get('category')

    if selected_cat_id:
        parameters = TestParameter.objects.filter(
            category_id=selected_cat_id
        ).select_related('group').order_by('group__order', 'order')
        for param in parameters:
            for old_key, old_val in old_fields.items():
                if (param.name.lower().strip() == old_key.lower().strip() or
                    param.name.lower() in old_key.lower() or
                    old_key.lower() in param.name.lower()):
                    prev_results[param.id] = {'result': '', 'standard': old_val}
                    break
        for group_key, group_items in groupby(parameters, key=lambda p: p.group):
            grouped_parameters.append({'group': group_key, 'params': list(group_items)})

    return render(request, 'coa/clone_from_old.html', {
        'form':               form,
        'old_coa':            old_coa,
        'old_fields':         old_fields,
        'categories':         categories,
        'grouped_parameters': grouped_parameters,
        'prev_results':       prev_results,
        'selected_cat_id':    selected_cat_id or '',
    })


def admin_required(view_func):
    from functools import wraps
    @wraps(view_func)
    @login_required
    def wrapper(request, *args, **kwargs):
        try:
            if request.user.is_superuser or request.user.profile.role == 'admin':
                return view_func(request, *args, **kwargs)
        except Exception:
            pass
        from django.http import HttpResponseForbidden
        return HttpResponseForbidden(
            "<h2>Access Denied</h2><p>Only admin users can manage accounts.</p>"
            "<a href='/coa/'>← Back</a>"
        )
    return wrapper


@admin_required
def user_list(request):
    users = User.objects.select_related('profile').order_by('username')
    return render(request, 'coa/user_list.html', {'users': users})


@admin_required
def user_create(request):
    if request.method == 'POST':
        username   = request.POST.get('username', '').strip()
        password   = request.POST.get('password', '').strip()
        first_name = request.POST.get('first_name', '').strip()
        last_name  = request.POST.get('last_name', '').strip()
        email      = request.POST.get('email', '').strip()
        role       = request.POST.get('role', 'analyst')

        if not username or not password:
            messages.error(request, 'Username and password are required.')
        elif User.objects.filter(username=username).exists():
            messages.error(request, f'Username "{username}" already exists.')
        else:
            user = User.objects.create_user(
                username=username, password=password,
                first_name=first_name, last_name=last_name, email=email,
            )
            UserProfile.objects.create(user=user, role=role, created_by=request.user)
            messages.success(request, f'User "{username}" created successfully.')
            return redirect('user_list')

    return render(request, 'coa/user_create.html', {
        'role_choices': UserProfile.ROLE_CHOICES,
    })


@admin_required
def user_edit(request, user_id):
    target_user = get_object_or_404(User, id=user_id)
    try:
        profile = target_user.profile
    except UserProfile.DoesNotExist:
        profile = UserProfile.objects.create(user=target_user, role='analyst')

    if request.method == 'POST':
        target_user.first_name = request.POST.get('first_name', '').strip()
        target_user.last_name  = request.POST.get('last_name', '').strip()
        target_user.email      = request.POST.get('email', '').strip()
        new_password = request.POST.get('password', '').strip()
        if new_password:
            target_user.set_password(new_password)
        target_user.save()

        profile.role = request.POST.get('role', profile.role)
        profile.save()
        messages.success(request, f'User "{target_user.username}" updated.')
        return redirect('user_list')

    return render(request, 'coa/user_edit.html', {
        'target_user':  target_user,
        'profile':      profile,
        'role_choices': UserProfile.ROLE_CHOICES,
    })


@admin_required
def user_delete(request, user_id):
    target_user = get_object_or_404(User, id=user_id)
    if target_user == request.user:
        messages.error(request, 'You cannot delete your own account.')
        return redirect('user_list')
    if request.method == 'POST':
        username = target_user.username
        target_user.delete()
        messages.success(request, f'User "{username}" deleted.')
        return redirect('user_list')
    return render(request, 'coa/user_confirm_delete.html', {'target_user': target_user})

@login_required
def download_coa_word(request, coa_id):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from io import BytesIO
    import os

    coa = get_object_or_404(COA, id=coa_id)
    results = coa.results.all().select_related(
        'parameter', 'parameter__group'
    ).order_by('parameter__group__order', 'parameter__order')
    grouped_results = []
    for group_key, group_items in groupby(results, key=lambda r: r.parameter.group):
        grouped_results.append({'group': group_key, 'items': list(group_items)})
    custom_fields  = list(coa.custom_fields.all())
    is_dry_extract = coa.category.is_dry_extract()

    static_dir = os.path.join(settings.BASE_DIR, 'coa', 'static', 'images')
    logo_path  = os.path.join(static_dir, 'logo.png')
    halal_path = os.path.join(static_dir, 'halal_badge.png')
    iso_path   = os.path.join(static_dir, 'iso_badge.png')
    gmp_path   = os.path.join(static_dir, 'gmp_badge.png')
    stamp_path = os.path.join(static_dir, 'stamp.png')

    doc = Document()
    section = doc.sections[0]
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def set_borders(cell, color='000000'):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'), 'single')
            tag.set(qn('w:sz'), '4')
            tag.set(qn('w:space'), '0')
            tag.set(qn('w:color'), color)
            tcBorders.append(tag)
        tcPr.append(tcBorders)

    def no_borders(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'), 'none')
            tag.set(qn('w:sz'), '0')
            tag.set(qn('w:color'), 'FFFFFF')
            tcBorders.append(tag)
        tcPr.append(tcBorders)

    # ── Header: Logo | Badges ──
    hdr = doc.add_table(rows=1, cols=2)
    lc = hdr.rows[0].cells[0]
    rc = hdr.rows[0].cells[1]
    no_borders(lc)
    no_borders(rc)
    if os.path.exists(logo_path):
        lc.paragraphs[0].add_run().add_picture(logo_path, height=Cm(1.5))
    rc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    br = rc.paragraphs[0].add_run()
    for bp in [halal_path, iso_path, gmp_path]:
        if os.path.exists(bp):
            br.add_picture(bp, height=Cm(1.2))

    # Divider
    d = doc.add_paragraph()
    dr = d.add_run('─' * 90)
    dr.font.size = Pt(7)

    # ── Title ──
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run('Certificate of Analysis')
    tr.bold = True
    tr.font.size = Pt(14)
    tr.underline = True

    # ── Product Info Table ──
    info = doc.add_table(rows=1, cols=2)
    info.style = 'Table Grid'
    lc = info.rows[0].cells[0]
    rc = info.rows[0].cells[1]
    lc.paragraphs[0].clear()
    rc.paragraphs[0].clear()

    def add_line(cell, label, value, italic=False):
        p = cell.add_paragraph()
        lb = p.add_run(f'{label} : ')
        lb.bold = True
        lb.font.size = Pt(9.5)
        vr = p.add_run(str(value) if value else '')
        vr.font.size = Pt(9.5)
        vr.font.italic = italic

    add_line(lc, 'Product Name', coa.product_name)
    add_line(lc, 'Batch No.', coa.batch_no)
    if coa.botanical_name:
        add_line(lc, 'Botanical Name', coa.botanical_name, italic=True)
    if coa.plant_part:
        add_line(lc, 'Part of Plant Used', coa.plant_part)
    add_line(rc, 'Mfg. Date', coa.manufacturing_date.strftime('%B-%Y') if coa.manufacturing_date else '—')
    add_line(rc, 'Exp. Date', coa.expiry_date.strftime('%B-%Y') if coa.expiry_date else '—')
    set_borders(lc)
    set_borders(rc)

    doc.add_paragraph()

    # ── Results Table ──
    col_count = 4 if is_dry_extract else 3
    rt = doc.add_table(rows=1, cols=col_count)
    rt.style = 'Table Grid'

    headers = ['Tests', 'Standard', 'Results']
    if is_dry_extract:
        headers.append('Reference')

    hdr_cells = rt.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], 'ECECEC')
        set_borders(hdr_cells[i])
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for section in grouped_results:
        if section['group']:
            gr = rt.add_row().cells
            for i in range(1, col_count):
                gr[0].merge(gr[i])
            set_cell_bg(gr[0], 'E2E2E2')
            set_borders(gr[0])
            run = gr[0].paragraphs[0].add_run(section['group'].name)
            run.bold = True
            run.font.size = Pt(9.5)

        for result in section['items']:
            row = rt.add_row().cells
            std = result.standard_override or result.parameter.specification or '—'
            data = [result.parameter.name, std, result.result or '—']
            if is_dry_extract:
                data.append(result.reference or '—')
            for i, val in enumerate(data):
                set_borders(row[i])
                run = row[i].paragraphs[0].add_run(val)
                run.font.size = Pt(9)
                if i > 0:
                    row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if custom_fields:
        for cf in custom_fields:
            if cf.is_heading:
                gr = rt.add_row().cells
                for i in range(1, col_count):
                    gr[0].merge(gr[i])
                set_cell_bg(gr[0], 'E2E2E2')
                set_borders(gr[0])
                run = gr[0].paragraphs[0].add_run(cf.field_name)
                run.bold = True
                run.font.size = Pt(9.5)
            else:
                row = rt.add_row().cells
                data = [cf.field_name, cf.specification or '—', cf.result or '—']
                if is_dry_extract:
                    data.append(cf.reference or '—')
                for i, val in enumerate(data):
                    set_borders(row[i])
                    run = row[i].paragraphs[0].add_run(val)
                    run.font.size = Pt(9)
                    if i > 0:
                        row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── Opinion ──
    op = doc.add_paragraph()
    opr = op.add_run('OPINION OF ANALYST: The above material complies with the prescribed API standards.')
    opr.bold = True
    opr.font.size = Pt(9)

    # ── Stamp ──
    if os.path.exists(stamp_path):
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.add_run().add_picture(stamp_path, height=Cm(1.8))

    doc.add_paragraph()

    # ── Signatures ──
    sig = doc.add_table(rows=1, cols=2)
    for cell, label in [(sig.rows[0].cells[0], 'Analysed By'), (sig.rows[0].cells[1], 'Approved By')]:
        no_borders(cell)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'\n\n{"_" * 30}\n{label}')
        run.bold = True
        run.font.size = Pt(9.5)

    doc.add_paragraph()

    # ── Note ──
    note = doc.add_paragraph()
    nr = note.add_run(
        'Note :- Since it is an herbal product, there is likely to be minor colour variation '
        'from batch to batch because of the seasonal variations of the raw materials. '
        'Colour does not affect the quality and efficacy of the product.'
    )
    nr.font.size = Pt(7)
    nr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # ── Company Footer ──
    cf_para = doc.add_paragraph()
    cf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cfr = cf_para.add_run(
        'Manufactured by: Hiya India Biotech Pvt. Ltd. & Supplied by: Nuplanet Ventures India Pvt. Ltd\n'
        'B-51, Okhla Industrial Area, Phase -1, New Delhi – 110020\n'
        'contact@rawble.com | marketing@rawble.com | admin@rawble.com'
    )
    cfr.font.size = Pt(7.5)
    cfr.font.color.rgb = RGBColor(0x1a, 0x5c, 0x1a)
    cfr.bold = True

    # ── Save ──
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_batch    = coa.batch_no.replace('/', '-')
    safe_product  = coa.product_name.replace(' ', '_').replace('/', '-')
    safe_customer = (coa.customer_name or 'General').replace(' ', '_').replace('/', '-')

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="COA_{safe_product}_{safe_customer}_{safe_batch}.docx"'
    return response